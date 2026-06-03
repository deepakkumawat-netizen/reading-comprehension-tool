import os
import sys
import json
import io
import re
from pathlib import Path
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from groq import Groq
from dotenv import load_dotenv

sys.path.insert(0, str(Path(__file__).parent))
from database import (init_db, create_session, save_comprehension,
                      get_session_history, get_all_comprehensions, save_rag_document)
from rag import rag_retriever
from nlp_adapter import get_grade_prompt_context, analyze_text_grade, GRADE_PROFILES, get_reading_counts
from mcp_tools import MCP_TOOLS, execute_mcp_tool
from security import (assert_public_url, read_upload_capped, client_ip,
                      generate_limiter, extract_limiter, upload_limiter)

load_dotenv()

GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_FALLBACK_MODELS = [
    GROQ_MODEL,
    "llama-3.1-8b-instant",
    "deepseek-r1-distill-llama-70b",
    "meta-llama/llama-4-scout-17b-16e-instruct",
]
_groq_client = None


def get_groq_client() -> Groq:
    global _groq_client
    if _groq_client is None:
        _groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    return _groq_client


def is_rate_limit_error(exc: Exception) -> bool:
    msg = str(exc).lower()
    return "429" in msg or "rate_limit" in msg or "rate limit" in msg or "tokens per day" in msg


@asynccontextmanager
async def lifespan(app: FastAPI):
    try:
        init_db()
    except Exception as e:
        print(f"[startup] DB init error: {e}")
    try:
        rag_retriever.build_index()
    except Exception as e:
        print(f"[startup] RAG build error: {e}")
    yield


app = FastAPI(title="Reading Comprehension Tool", version="1.0.0", lifespan=lifespan)

# Same-origin SPA + localhost dev only — used to be `["*"]`, which let any
# origin POST to the generate endpoint and drain the Groq quota. Override
# via ALLOWED_ORIGINS env var if you need to embed elsewhere.
_default_origins = "https://reading-comprehension-tool.onrender.com,http://localhost:5173,http://localhost:5174,http://127.0.0.1:5173"
ALLOWED_ORIGINS = [o.strip() for o in os.getenv("ALLOWED_ORIGINS", _default_origins).split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=False,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)


# ── Cache-Control middleware ─────────────────────────────────────────────────
# The SPA's index.html references hashed asset filenames (index-{hash}.js).
# When we deploy a new build, the HTML changes to reference a NEW hash but
# browsers with cached HTML keep loading the OLD JS bundle until their
# cache expires — which is exactly the "user reports the same bug after we
# fix it" loop. Serve HTML with no-cache, hashed assets with long-immutable.
@app.middleware("http")
async def cache_control_headers(request, call_next):
    response = await call_next(request)
    path = request.url.path
    is_html = path == "/" or path.endswith(".html")
    is_hashed_asset = (
        path.startswith("/assets/")
        and any(path.endswith(ext) for ext in (".js", ".css", ".woff", ".woff2", ".ttf"))
    )
    if is_html:
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    elif is_hashed_asset:
        response.headers["Cache-Control"] = "public, max-age=31536000, immutable"
    return response


class ComprehensionRequest(BaseModel):
    topic: str
    grade_level: int
    learning_objective: str
    board: Optional[str] = None  # CBSE, ICSE, RBSE, etc. — woven into the prompt for curriculum alignment
    source_text: Optional[str] = None
    additional_context: Optional[str] = None
    session_id: Optional[str] = None


class SessionCreate(BaseModel):
    metadata: Optional[dict] = None


class MCPToolCall(BaseModel):
    tool_name: str
    arguments: dict


class RAGDocRequest(BaseModel):
    content: str
    topic: Optional[str] = ""
    grade_level: Optional[int] = 0


class CompleteAnswerRequest(BaseModel):
    question: str
    passage_text: str
    grade_level: int
    word_limit: int = 35
    question_type: Optional[str] = "literal"
    answer_hint: Optional[str] = ""


# ── Sessions ──────────────────────────────────────────────────────────────────

@app.post("/api/sessions")
async def new_session(req: SessionCreate):
    session_id = create_session(req.metadata)
    return {"session_id": session_id}


@app.get("/api/sessions/{session_id}/history")
async def session_history(session_id: str):
    return {"session_id": session_id, "history": get_session_history(session_id)}


@app.get("/api/comprehensions")
async def list_comprehensions(limit: int = 20):
    return {"comprehensions": get_all_comprehensions(limit)}


# ── Complete Answer (NLP-calibrated per grade) ────────────────────────────────

@app.post("/api/reading/complete-answer")
async def complete_answer(req: CompleteAnswerRequest):
    p = GRADE_PROFILES.get(req.grade_level, GRADE_PROFILES[7])
    grade_ctx = get_grade_prompt_context(req.grade_level)

    prompt = f"""You are an expert educator writing a model answer for a Grade {req.grade_level} student.

{grade_ctx}

READING PASSAGE:
{req.passage_text[:2000]}

QUESTION TYPE: {req.question_type}
QUESTION: {req.question}

STRICT WORD LIMIT: {req.word_limit} words maximum. Count every word — do NOT exceed this limit.

TASK: Write a model answer in EXACTLY {req.word_limit} words or fewer.
RULES:
1. HARD LIMIT: Your answer must be {req.word_limit} words or fewer. This is non-negotiable.
2. Use ONLY Grade {req.grade_level} vocabulary: {p['vocab']}
3. Sentence structure for Grade {req.grade_level}: {p['sentence']}
4. Cognitive level: {p['blooms']}
5. Reference the passage text as evidence but stay within the word limit.
6. Write ONLY the answer — no "Answer:", no labels, no explanation outside the answer.

Answer ({req.word_limit} words max):"""

    last_error = ""
    for model in GROQ_FALLBACK_MODELS:
        try:
            response = get_groq_client().chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=min(req.word_limit * 3, 300),
            )
            # Groq can return content=None when the model produces only a
            # refusal/tool message. Calling .strip() on None raises
            # AttributeError, which was previously swallowed by the bare
            # `except Exception` below, ate the fallback chain, and surfaced
            # as a generic "Failed after N attempts" to the user. Coerce to
            # empty string so we move to the next model cleanly.
            raw_content = (response.choices[0].message.content if response.choices else None) or ""
            answer = raw_content.strip()
            answer = re.sub(r'^(answer|model answer|response)\s*:\s*', '', answer, flags=re.IGNORECASE).strip()
            if not answer:
                # Skip empty/refusal completions and try the next model.
                last_error = f"empty completion from {model}"
                continue

            # Hard truncate to word limit as safety net
            words = answer.split()
            if len(words) > req.word_limit:
                answer = " ".join(words[:req.word_limit])
                # End at last complete sentence if possible
                for punct in ('.', '!', '?'):
                    last = answer.rfind(punct)
                    if last > len(answer) // 2:
                        answer = answer[:last + 1]
                        break

            return {"answer": answer}
        except Exception as exc:
            last_error = str(exc)
            continue

    raise HTTPException(status_code=503, detail=f"Failed after {len(GROQ_FALLBACK_MODELS)} attempts: {last_error}")


# ── Generate ──────────────────────────────────────────────────────────────────

def _count_syllables(word: str) -> int:
    """Rough syllable counter — vowel groups, minus silent terminal 'e'."""
    word = (word or "").lower().strip(".,!?;:'\"")
    if not word:
        return 1
    count = len(re.findall(r'[aeiouy]+', word))
    if word.endswith('e') and count > 1:
        count -= 1
    return max(count, 1)


# Per-grade complexity caps for the vocabulary_in_context items. The
# GRADE_PROFILES are instructive in the prompt but the model doesn't
# always follow them at Grade 1-3 — it will happily pull
# "photosynthesis" from the source. Only applied to Grade 1-3.
GRADE_VOCAB_CAPS = {
    1: {"max_syllables": 2, "max_chars": 7},
    2: {"max_syllables": 3, "max_chars": 8},
    3: {"max_syllables": 3, "max_chars": 9},
}


def _check_grade_complexity(data: dict, grade_level: int) -> "str | None":
    """For Grade 1-3, reject vocabulary_in_context items whose target word
    exceeds the syllable/length cap for the grade. Returns the error string
    so the retry prompt can call out which words to swap."""
    caps = GRADE_VOCAB_CAPS.get(grade_level)
    if not caps:
        return None
    vic = data.get("vocabulary_in_context", {})
    too_complex = []
    for item in vic.get("items", []) or []:
        word = (item.get("word") or "").strip().lower()
        if not word:
            continue
        if _count_syllables(word) > caps["max_syllables"] or len(word) > caps["max_chars"]:
            too_complex.append(word)
    if too_complex:
        return (
            f"Grade {grade_level} vocabulary must be at most {caps['max_syllables']} syllables and "
            f"{caps['max_chars']} letters per word. These words are too complex: "
            f"{', '.join(too_complex[:5])}. REPLACE each with a simpler Grade {grade_level} word "
            f"on the same topic (1-2 syllable sight word, CVC pattern, decodable phonics)."
        )
    return None


def _validate_reading(data: dict, grade_level: int = 7) -> "str | None":
    """Return an error description string if the comprehension JSON is invalid, else None.
    Validates structure AND that the passage length roughly matches the grade."""
    byr = data.get("before_you_read")
    if not isinstance(byr, dict) or not isinstance(byr.get("questions"), list) or len(byr["questions"]) < 1:
        return "before_you_read.questions is missing or empty"

    passage = data.get("passage")
    if not isinstance(passage, dict) or not passage.get("text"):
        return "passage.text is missing or empty"

    # Reject passages grossly longer than the grade's range (e.g. Grade 1 getting
    # 263 words). Uses 1.6x the grade max as the ceiling to allow some flexibility.
    try:
        p = GRADE_PROFILES.get(grade_level, GRADE_PROFILES[7])
        rng = p.get("passage_words", "")
        nums = [int(x) for x in re.findall(r"\d+", rng)]
        if nums:
            max_words = max(nums)
            actual = len(passage["text"].split())
            ceiling = int(max_words * 1.6)
            if actual > ceiling:
                return (f"passage is {actual} words but Grade {grade_level} must be "
                        f"{rng} words. Rewrite it MUCH shorter — no more than {max_words} words.")
    except Exception:
        pass

    tdq = data.get("text_dependent_questions")
    if not isinstance(tdq, dict) or not isinstance(tdq.get("questions"), list) or len(tdq["questions"]) < 1:
        return "text_dependent_questions.questions is missing or empty"

    vic = data.get("vocabulary_in_context")
    if not isinstance(vic, dict) or not isinstance(vic.get("items"), list) or len(vic["items"]) < 1:
        return "vocabulary_in_context.items is missing or empty"

    return None


@app.post("/api/reading/generate")
async def generate_comprehension(req: ComprehensionRequest, request: Request):
    # Per-IP rate limit on the most-expensive endpoint. Without this, anyone
    # who finds the Render URL can drain the daily Groq quota in a loop.
    await generate_limiter.check(client_ip(request))
    session_id = req.session_id or create_session()

    rag_retriever.build_index()
    rag_context = rag_retriever.build_context(
        f"{req.topic} grade {req.grade_level} {req.learning_objective}",
        grade_level=req.grade_level
    )

    grade_ctx = get_grade_prompt_context(req.grade_level)

    # Pre-compute optional blocks to avoid backslashes inside f-string expressions
    source_block = (
        "\nSOURCE MATERIAL (MANDATORY — base the passage's facts, names, examples and "
        "vocabulary on THIS content only; do not invent facts that contradict it):\n"
        f"---\n{req.source_text[:6000]}\n---\n"
    ) if req.source_text else ""
    additional_block = f"Additional Context: {req.additional_context}" if req.additional_context else ""
    rag_block = f"\n{rag_context}" if rag_context else ""
    # Curriculum board (CBSE / ICSE / RBSE / ...) becomes a hard constraint
    # in the prompt — the passage, vocabulary register, and question style
    # must match this board's published syllabus for the selected grade.
    board_block = (
        f"\nCURRICULUM BOARD: {req.board}\n"
        f"The passage, vocabulary, and questions must align with the {req.board} syllabus and "
        f"exam style for Grade {req.grade_level}. Use vocabulary, examples, and context that an "
        f"actual {req.board} Grade {req.grade_level} student would encounter in their official "
        f"textbooks and exams.\n"
    ) if req.board else ""
    ctx_block = f"{board_block}{source_block}{additional_block}\n{rag_block}".strip()

    def _build_prompt(extra_instructions: str = "") -> str:
        p = GRADE_PROFILES.get(req.grade_level, GRADE_PROFILES[7])
        word_range = p["passage_words"]
        c = get_reading_counts(req.grade_level)

        low_grade_block = f"\n- Passage must be {word_range} words.\n"

        return f"""You are an expert reading specialist and curriculum designer.
Your task is to create a complete, grade-calibrated Reading Comprehension activity.

{grade_ctx}
{low_grade_block}
CONTENT DETAILS:
Topic: {req.topic}
Learning Objective: {req.learning_objective}
{ctx_block}

CRITICAL RULES:
1. Every word you write — passage, questions, instructions, hints — must match Grade {req.grade_level} level EXACTLY.
2. The passage MUST be {word_range} words — count carefully.
3. Generate EXACTLY {c['total_q']} text-dependent questions (calibrated for Grade {req.grade_level} attention) and EXACTLY {c['vocab']} vocabulary items. Do NOT write all literal questions.
4. Vocabulary in Context words must come directly from the passage.
5. Before You Read questions must activate prior knowledge at a Grade {req.grade_level} cognitive level.
{"6. SOURCE MATERIAL is provided above and is the AUTHORITATIVE basis for this passage. The passage MUST be a grade-level rewrite/summary of the SOURCE MATERIAL — every fact, name, number, date, event, term and example must come directly from it. Do NOT invent content from your own knowledge if it contradicts or is absent from the source. Text-dependent questions must reference the rewritten passage (which reflects the source); Vocabulary in Context words must be picked from words actually present in the source." if req.source_text else ""}
{extra_instructions}

Return ONLY valid JSON. No markdown fences. No prose outside the JSON.

{{
  "before_you_read": {{
    "title": "Before You Read",
    "instructions": "Grade {req.grade_level}-appropriate activation prompt here (1 sentence).",
    "questions": [
      {{"number": 1, "question": "Grade {req.grade_level} prior-knowledge question about {req.topic}", "type": "activation"}},
      {{"number": 2, "question": "Grade {req.grade_level} prediction question about the passage", "type": "prediction"}},
      {{"number": 3, "question": "Grade {req.grade_level} inquiry question the student wonders about {req.topic}", "type": "inquiry"}}
    ]
  }},
  "annotation_guide": {{
    "title": "Annotation Guide",
    "instructions": "Grade {req.grade_level}-appropriate reading strategy instruction (1 sentence).",
    "symbols": [
      {{"symbol": "⭐", "meaning": "Grade {req.grade_level} explanation of main idea marking"}},
      {{"symbol": "?", "meaning": "Grade {req.grade_level} explanation of confusion marking"}},
      {{"symbol": "!", "meaning": "Grade {req.grade_level} explanation of interesting info marking"}},
      {{"symbol": "→", "meaning": "Grade {req.grade_level} explanation of cause-effect marking"}},
      {{"symbol": "circle", "meaning": "Grade {req.grade_level} explanation of vocabulary marking"}}
    ]
  }},
  "passage": {{
    "title": "Engaging title relevant to {req.topic}",
    "text": "Write the FULL passage here. Must be {word_range} words. Use paragraph breaks (\\n\\n). Every sentence must match Grade {req.grade_level} syntax and vocabulary.",
    "word_count": "actual number of words in the passage you wrote"
  }},
  "text_dependent_questions": {{
    "title": "Text-Dependent Questions",
    "instructions": "Grade {req.grade_level}-appropriate instruction for answering with text evidence.",
    "questions": [
      {{"number": 1, "question": "Question at Grade {req.grade_level} level", "type": "literal", "answer_hint": "Paragraph evidence"}},
      ... Generate EXACTLY {c['total_q']} questions total for Grade {req.grade_level}:
          {c['literal']} literal (type "literal"), {c['inferential']} inferential (type "inferential"){', ' + str(c['higher']) + ' higher-order Analyze/Evaluate (type "critical_thinking")' if c['higher'] else ''}.
          Number them 1..{c['total_q']}. Each needs an answer_hint pointing to passage evidence.
    ]
  }},
  "vocabulary_in_context": {{
    "title": "Vocabulary in Context",
    "instructions": "Grade {req.grade_level}-appropriate vocabulary strategy instruction.",
    "items": [
      {{
        "word": "actual word from the passage appropriate for Grade {req.grade_level}",
        "sentence_from_passage": "Copy the exact sentence from your passage containing this word.",
        "context_clue_type": "definition|example|contrast|inference",
        "activity": "Grade {req.grade_level}-appropriate activity using this word",
        "answer": "Grade {req.grade_level}-appropriate answer"
      }},
      ... EXACTLY {c['vocab']} items total, each word from the passage
    ]
  }}
}}"""

    def _sse(obj: dict) -> str:
        return f"data: {json.dumps(obj)}\n\n"

    def stream_gen():
        max_attempts = 5
        extra_instructions = ""
        last_reason = ""
        model_idx = 0

        for attempt in range(1, max_attempts + 1):
            if attempt > 1:
                yield _sse({"type": "retry", "attempt": attempt, "reason": last_reason})

            current_model = GROQ_FALLBACK_MODELS[min(model_idx, len(GROQ_FALLBACK_MODELS) - 1)]
            yield _sse({"type": "progress", "message": f"Attempt {attempt}: calling {current_model}…"})

            prompt = _build_prompt(extra_instructions)
            collected_chunks = []

            try:
                stream = get_groq_client().chat.completions.create(
                    model=current_model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.75,
                    max_tokens=4500,
                    stream=True,
                )

                for chunk in stream:
                    # Groq sometimes emits a usage-only / keep-alive chunk
                    # with empty choices — indexing [0] then raises
                    # IndexError, which the outer except swallows and
                    # triggers a spurious "Model error — switching to…"
                    # mid-success. Skip those frames silently.
                    if not chunk.choices:
                        continue
                    delta = chunk.choices[0].delta.content or ""
                    if delta:
                        collected_chunks.append(delta)
                        yield _sse({"type": "token", "content": delta})

            except Exception as exc:
                last_reason = str(exc)
                if model_idx < len(GROQ_FALLBACK_MODELS) - 1:
                    model_idx += 1
                    next_model = GROQ_FALLBACK_MODELS[model_idx]
                    yield _sse({"type": "status", "message": f"Model error — switching to {next_model}…"})
                    extra_instructions = ""
                else:
                    extra_instructions = f"IMPORTANT: Fix the following error from the previous attempt: {last_reason}\n"
                continue

            raw = "".join(collected_chunks).strip()
            for fence in ("```json", "```"):
                if raw.startswith(fence):
                    raw = raw[len(fence):]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()
            # Remove control characters invalid inside JSON strings (keep \t \n \r)
            import re as _re
            raw = _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', raw)

            yield _sse({"type": "status", "message": "Parsing JSON response…"})

            # Isolate the JSON object (drop any stray prose before/after)
            first, last = raw.find("{"), raw.rfind("}")
            if first != -1 and last != -1 and last > first:
                raw = raw[first:last + 1]

            data = None
            try:
                data = json.loads(raw)
            except json.JSONDecodeError as exc:
                # LLMs often emit unescaped quotes/apostrophes or trailing commas.
                # json-repair fixes these instead of forcing another full retry.
                try:
                    from json_repair import repair_json
                    repaired = repair_json(raw)
                    data = json.loads(repaired)
                except Exception:
                    last_reason = f"Invalid JSON: {exc}"
                    extra_instructions = (
                        "CRITICAL: Your previous response was not valid JSON. "
                        "Return ONLY a raw JSON object — no markdown fences, no prose. "
                        "Escape every double-quote inside string values as \\\".\n"
                    )
                    # Rotate to the next model on JSON-parse failures too. The
                    # old code only advanced model_idx on hard exceptions, so
                    # if model 0 reliably emitted malformed JSON for a given
                    # prompt, all 5 attempts hit model 0 with the same broken
                    # pattern. Advancing here gives the fallback list a real
                    # chance to recover.
                    if model_idx < len(GROQ_FALLBACK_MODELS) - 1:
                        model_idx += 1
                        yield _sse({"type": "status", "message": f"Switching to {GROQ_FALLBACK_MODELS[model_idx]}…"})
                    continue

            yield _sse({"type": "status", "message": "Validating comprehension structure…"})

            validation_error = _validate_reading(data, req.grade_level)
            if validation_error:
                last_reason = f"Validation failed: {validation_error}"
                extra_instructions = (
                    f"IMPORTANT: Fix this validation error from your previous attempt: {validation_error}. "
                    "Ensure before_you_read has >=3 questions, passage.text is present, "
                    f"passage is {GRADE_PROFILES.get(req.grade_level, GRADE_PROFILES[7])['passage_words']} words, "
                    "text_dependent_questions has >=6 questions, and vocabulary_in_context has >=5 items.\n"
                )
                # Same rotation as the JSON-parse failure path above. A model
                # that overshoots passage length once with temperature 0.75
                # tends to overshoot it again with the same prompt — try a
                # different model on the next attempt instead of retrying.
                if model_idx < len(GROQ_FALLBACK_MODELS) - 1:
                    model_idx += 1
                    yield _sse({"type": "status", "message": f"Switching to {GROQ_FALLBACK_MODELS[model_idx]}…"})
                continue

            # For Grade 1-3, also enforce that vocabulary_in_context words
            # actually match grade complexity. The model often pulls a complex
            # source word like "photosynthesis" into a Grade 1 worksheet even
            # with the NLP calibration in the prompt — this is the
            # post-generation enforcement step.
            complexity_error = _check_grade_complexity(data, req.grade_level)
            if complexity_error:
                last_reason = f"Grade-complexity failed: {complexity_error}"
                yield _sse({"type": "status", "message": "Words too complex for the grade — regenerating with simpler vocabulary…"})
                extra_instructions = complexity_error + "\n"
                continue

            # Annotate passage with readability metrics
            passage_text = data.get("passage", {}).get("text", "")
            readability = analyze_text_grade(passage_text)
            if readability:
                data["passage"]["readability"] = readability
                word_count = readability.get("word_count", 0)
                if word_count:
                    data["passage"]["word_count"] = word_count


            yield _sse({"type": "status", "message": "Saving comprehension…"})

            full_content = {**data, "rag_context_used": bool(rag_context)}

            try:
                comp_id = save_comprehension(
                    session_id=session_id,
                    topic=req.topic,
                    grade_level=req.grade_level,
                    learning_objective=req.learning_objective,
                    content=full_content,
                )

                save_rag_document(
                    content=(
                        f"reading comprehension topic {req.topic} grade {req.grade_level} "
                        f"objective {req.learning_objective} passage: {passage_text[:300]}"
                    ),
                    doc_type="comprehension",
                    topic=req.topic,
                    grade_level=req.grade_level,
                )
                rag_retriever.build_index()
            except Exception as exc:
                yield _sse({"type": "error", "message": f"Database error: {exc}"})
                return

            yield _sse({
                "type": "complete",
                "session_id": session_id,
                "comprehension_id": comp_id,
                "comprehension": full_content,
            })
            return

        # Exhausted all retries
        yield _sse({"type": "error", "message": f"Failed after {max_attempts} attempts. Last error: {last_reason}"})

    return StreamingResponse(
        stream_gen(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


# ── Export DOCX ───────────────────────────────────────────────────────────────

@app.post("/api/reading/export/docx")
async def export_docx(payload: dict):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    comp = payload.get("comprehension", {})
    topic = payload.get("topic", "Reading")
    grade = payload.get("grade_level", "")
    objective = payload.get("learning_objective", "")

    doc = Document()

    title = doc.add_heading("Reading Comprehension Activity", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Topic: {topic}  |  Grade: {grade}  |  Objective: {objective}")
    doc.add_paragraph("Name: ____________________________   Date: _______________")
    doc.add_paragraph()

    # All field accesses below use .get() with sensible fallbacks. The LLM
    # validator only checks counts, not per-item keys, so any item with a
    # missing field would otherwise raise KeyError → 500 → the frontend
    # downloaded an "{detail: ...}" JSON blob saved as .docx (which Word
    # then reports as a corrupt file).

    # Before You Read
    byr = comp.get("before_you_read", {})
    if byr:
        doc.add_heading(byr.get("title", "Before You Read"), 1)
        doc.add_paragraph(byr.get("instructions", ""))
        for i, q in enumerate(byr.get("questions", []), 1):
            num = q.get("number", i)
            doc.add_paragraph(f"{num}. {q.get('question', '')}")
            doc.add_paragraph("   Answer: ____________________________________________")
        doc.add_paragraph()

    # Annotation Guide
    ag = comp.get("annotation_guide", {})
    if ag:
        doc.add_heading(ag.get("title", "Annotation Guide"), 1)
        doc.add_paragraph(ag.get("instructions", ""))
        for s in ag.get("symbols", []):
            doc.add_paragraph(f"  {s.get('symbol', '')} = {s.get('meaning', '')}")
        doc.add_paragraph()

    # Passage
    passage = comp.get("passage", {})
    if passage:
        doc.add_heading(passage.get("title", "Reading Passage"), 1)
        for para in passage.get("text", "").split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.add_paragraph()

    # Text-Dependent Questions
    tdq = comp.get("text_dependent_questions", {})
    if tdq:
        doc.add_heading(tdq.get("title", "Text-Dependent Questions"), 1)
        doc.add_paragraph(tdq.get("instructions", ""))
        for i, q in enumerate(tdq.get("questions", []), 1):
            num = q.get("number", i)
            doc.add_paragraph(f"{num}. {q.get('question', '')}")
            doc.add_paragraph("   Answer: ____________________________________________")
            doc.add_paragraph("   ____________________________________________________")
        doc.add_paragraph()

    # Vocabulary in Context
    vic = comp.get("vocabulary_in_context", {})
    if vic:
        doc.add_heading(vic.get("title", "Vocabulary in Context"), 1)
        doc.add_paragraph(vic.get("instructions", ""))
        for i, item in enumerate(vic.get("items", []), 1):
            doc.add_paragraph(f"{i}. Word: \"{item.get('word', '')}\"")
            doc.add_paragraph(f"   From the text: \"{item.get('sentence_from_passage', '')}\"")
            doc.add_paragraph(f"   {item.get('activity', '')}")
            doc.add_paragraph("   My answer: _________________________________________")
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="reading_{topic}.docx"'},
    )


# ── RAG document upload ───────────────────────────────────────────────────────

@app.post("/api/rag/add-text")
async def add_rag_text(req: RAGDocRequest):
    doc_id = save_rag_document(req.content, "knowledge", req.topic, req.grade_level)
    rag_retriever.build_index()
    return {"success": True, "doc_id": doc_id}


@app.post("/api/rag/add-file")
async def add_rag_file(request: Request, file: UploadFile = File(...)):
    await upload_limiter.check(client_ip(request))
    # Stream-read with a 10MB cap. Previously `await file.read()` buffered
    # the entire upload — a multi-GB POST would OOM the container.
    raw = await read_upload_capped(file)
    content = ""
    if file.filename.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(raw))
        content = " ".join(p.extract_text() or "" for p in reader.pages)
    elif file.filename.endswith(".docx"):
        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(raw))
        content = " ".join(p.text for p in doc.paragraphs)
    else:
        content = raw.decode("utf-8", errors="ignore")

    content = (content or "").strip()
    doc_id = save_rag_document(content[:6000], "file", file.filename, 0)
    rag_retriever.build_index()
    # Return the extracted text so the frontend can pass it as source_text
    # for the generator — otherwise the model ignored uploaded documents.
    return {
        "success": True,
        "doc_id": doc_id,
        "chars_indexed": len(content),
        "text": content[:8000],
        "filename": file.filename,
    }


# ── Hero image proxy (NanoBanana AI) ─────────────────────────────────────────
# Frontend's Landing page asks for `/api/hero-image?seed=N` and gets back a
# fresh cartoon JPEG generated by NanoBanana. The secret key lives in
# NANOBANANA_API_KEY on the server, never reaches the browser.
#
# NanoBanana is async — POST a generate task, poll record-info until done,
# then fetch the result image. End-to-end takes ~15s.

READING_HERO_PROMPTS = [
    "3D Pixar cartoon of a happy child reading an open storybook with characters and scenes coming alive out of the pages, magical golden glow, bright colors, clean white background",
    "3D Pixar cartoon of a cozy reading nook with a kid in pajamas reading on a beanbag, surrounded by stacks of colorful books and a sleeping cat, clean white background",
    "3D Pixar cartoon of a smiling student wearing reading glasses holding a book with question marks and lightbulbs floating around their head, classroom setting, clean white background",
    "3D Pixar cartoon of a magical library with floating books opening their pages mid-air and golden text streaming out, bright cheerful colors, clean white background",
    "3D Pixar cartoon of three diverse kids sitting in a circle each reading different colorful books, friendly smiles, clean white background, educational scene",
    "3D Pixar cartoon of an open book with a paper boat sailing out of it on a wave of words and stars, dreamy magical scene, bright vibrant colors, clean white background",
]


@app.get("/api/hero-image")
async def hero_image(request: Request, seed: Optional[int] = None):
    """Proxy to Hugging Face FLUX.1-schnell — fresh cartoon hero image per
    visit. HF token in HF_API_TOKEN env var; ~5-10s warm, up to ~60s cold."""
    await extract_limiter.check(client_ip(request))
    hf_token = (os.getenv("HF_API_TOKEN") or "").strip()
    if not hf_token:
        raise HTTPException(status_code=503, detail="HF_API_TOKEN not configured")

    import random as _random
    if seed is None:
        seed = _random.randint(1, 999999)
    prompt = READING_HERO_PROMPTS[seed % len(READING_HERO_PROMPTS)]

    try:
        import httpx
        async with httpx.AsyncClient(timeout=120.0) as cx:
            r = await cx.post(
                "https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-schnell",
                headers={
                    "Authorization": f"Bearer {hf_token}",
                    "Content-Type": "application/json",
                    "x-wait-for-model": "true",
                    "Accept": "image/png",
                },
                json={"inputs": prompt, "parameters": {"seed": seed, "num_inference_steps": 4}},
            )
        if r.status_code == 503:
            raise HTTPException(status_code=503, detail="HF model is warming up; try again in 30s")
        if r.status_code == 429:
            raise HTTPException(status_code=429, detail="HF free-tier rate limit hit")
        if r.status_code != 200:
            detail = r.text[:200] if r.text else f"HTTP {r.status_code}"
            raise HTTPException(status_code=502, detail=f"HF returned {r.status_code}: {detail}")
        from fastapi.responses import Response
        return Response(
            content=r.content,
            media_type=r.headers.get("content-type", "image/png"),
            headers={"Cache-Control": "public, max-age=3600"},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Hero image generation failed: {e}")


@app.post("/api/extract-url")
async def extract_url(req: dict, request: Request):
    """Fetch a webpage and return the cleaned article text as source_text."""
    await extract_limiter.check(client_ip(request))
    # Block private / loopback / metadata IPs so this endpoint can't be used
    # to read AWS/Render IMDS or pivot into the internal network.
    url = assert_public_url((req.get("url") or "").strip())
    try:
        import httpx
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(follow_redirects=True, timeout=20.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ReadingTool/1.0)"}) as cx:
            r = await cx.get(url)
            # Re-validate the final URL after redirects so a public URL can't
            # 302 us into a private host.
            assert_public_url(str(r.url))
            r.raise_for_status()
            soup = BeautifulSoup(r.text, "html.parser")
            for bad in soup(["script", "style", "noscript", "iframe", "nav", "footer", "header", "form", "aside"]):
                bad.decompose()
            title = (soup.title.string or "").strip() if soup.title else ""
            main = soup.find("main") or soup.find("article") or soup.body or soup
            text = re.sub(r"\s+\n", "\n", main.get_text("\n", strip=True))
            text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if not text:
            raise HTTPException(status_code=422, detail="Could not extract readable text from this page.")
        return {"success": True, "title": title, "url": url, "text": text[:8000], "chars": len(text)}
    except HTTPException:
        raise
    except httpx.HTTPStatusError as e:
        # Surface upstream 4xx as 4xx — previously a 404 page came back as
        # "500 URL fetch failed: 404 Not Found".
        raise HTTPException(status_code=e.response.status_code,
                            detail=f"URL fetch failed: {e.response.status_code} {e.response.reason_phrase}")
    except httpx.RequestError as e:
        raise HTTPException(status_code=502, detail=f"Could not reach URL: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"URL fetch failed: {e}")


@app.post("/api/auto-fields")
async def auto_fields(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    """Read uploaded text and propose a Topic + Learning Objective + grade.

    Saves the teacher from typing — they upload a PDF / URL / YouTube
    transcript and we suggest what to put in the Topic and Learning
    Objective fields, calibrated for the grade if they've picked one.
    """
    source_text = (req.get("source_text") or "").strip()
    if not source_text:
        raise HTTPException(status_code=400, detail="source_text is required.")
    grade = req.get("grade_level")
    board = (req.get("board") or "").strip()
    grade_hint = f"Calibrate the topic + objective for Grade {grade} students. " if grade else ""
    board_hint = f"This is for the {board} curriculum — phrase the topic and objective to match {board} terminology and exam style. " if board else ""

    prompt = (
        "You are an expert reading-comprehension curriculum designer.\n"
        "Read the SOURCE MATERIAL below and propose a clean, classroom-ready\n"
        "  • Topic (a short noun phrase, 4-10 words, naming the main idea)\n"
        "  • Learning Objective (one sentence starting with 'Students will…' "
        "describing what the student will be able to do after reading).\n"
        f"{grade_hint}{board_hint}"
        "Return ONLY a JSON object with keys 'topic' and 'learning_objective'. "
        "No markdown, no prose outside the JSON.\n\n"
        "SOURCE MATERIAL:\n---\n"
        f"{source_text[:6000]}\n---\n\n"
        'JSON: {"topic": "...", "learning_objective": "Students will ..."}'
    )

    try:
        client = get_groq_client()
        completion = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You return strict JSON. No markdown fences."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=400,
            response_format={"type": "json_object"},
        )
        import json as _json
        raw = completion.choices[0].message.content.strip()
        data = _json.loads(raw)
        topic = (data.get("topic") or "").strip()
        objective = (data.get("learning_objective") or "").strip()
        if not topic and not objective:
            raise HTTPException(status_code=502, detail="Model returned no fields.")
        return {"success": True, "topic": topic, "learning_objective": objective}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"auto-fields failed: {e}")


async def _youtube_metadata_fallback(video_id: str, url: str) -> dict | None:
    """When the transcript API is IP-blocked, build source_text from whatever
    public metadata we can grab. Tries three sources in order so we always
    return *something* the LLM can use as the basis for content:

      1. YouTube oEmbed (always works, gives title + channel — any IP)
      2. noembed.com aggregator (often includes a longer description)
      3. youtube.com/watch with consent cookie (full shortDescription if
         the cloud IP isn't behind a consent wall)
    """
    import httpx
    title = ""
    author = ""
    description = ""

    # 1. oEmbed — guaranteed title + author from any IP.
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=10.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ReadingTool/1.0)"}) as cx:
            r = await cx.get("https://www.youtube.com/oembed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}", "format": "json"})
            if r.status_code == 200:
                j = r.json()
                title = (j.get("title") or "").strip()
                author = (j.get("author_name") or "").strip()
    except Exception as _e:
        print(f"[YouTube oEmbed] failed: {_e}")

    # 2. noembed.com — third-party aggregator that often surfaces the description.
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=10.0,
            headers={"User-Agent": "Mozilla/5.0 (compatible; ReadingTool/1.0)"}) as cx:
            r = await cx.get("https://noembed.com/embed",
                params={"url": f"https://www.youtube.com/watch?v={video_id}"})
            if r.status_code == 200:
                j = r.json()
                if not title:  title  = (j.get("title") or "").strip()
                if not author: author = (j.get("author_name") or "").strip()
                description = (j.get("description") or "").strip()
    except Exception as _e:
        print(f"[YouTube noembed] failed: {_e}")

    # 3. Direct watch-page scrape — passes a consent cookie to bypass the EU wall.
    try:
        from bs4 import BeautifulSoup
        async with httpx.AsyncClient(follow_redirects=True, timeout=15.0,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                              "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept-Language": "en-US,en;q=0.9",
                "Cookie": "CONSENT=YES+cb.20210328-17-p0.en+FX+000",
            }) as cx:
            r = await cx.get(f"https://www.youtube.com/watch?v={video_id}")
            if r.status_code == 200:
                soup = BeautifulSoup(r.text, "html.parser")
                if not title:
                    ot = soup.find("meta", attrs={"property": "og:title"})
                    if ot: title = (ot.get("content") or "").strip()
                ot_desc = soup.find("meta", attrs={"property": "og:description"})
                if ot_desc and not description:
                    description = (ot_desc.get("content") or "").strip()
                for s in soup.find_all("script"):
                    txt = s.string or ""
                    if "shortDescription" in txt:
                        mm = re.search(r'"shortDescription":"((?:\\.|[^"\\])*)"', txt)
                        if mm:
                            full = mm.group(1).encode("utf-8").decode("unicode_escape", errors="ignore")
                            if len(full) > len(description):
                                description = full
                            break
    except Exception as _e:
        print(f"[YouTube watch-page] failed: {_e}")

    pieces = []
    if title:       pieces.append(f"Video title: {title}")
    if author:      pieces.append(f"Channel: {author}")
    if description: pieces.append(f"\n{description}")
    text = "\n".join(pieces).strip()
    if not text or len(text) < 20:
        return None
    return {"title": title or f"YouTube video {video_id}", "text": text}


@app.post("/api/extract-youtube")
async def extract_youtube(req: dict, request: Request):
    await extract_limiter.check(client_ip(request))
    """Fetch a YouTube transcript and return it as source_text. If YouTube
    blocks the transcript API (typical on cloud IPs), fall back to scraping
    the video's title + description so we still have usable source text."""
    url = (req.get("url") or "").strip()
    if not url:
        raise HTTPException(status_code=400, detail="url is required")
    m = re.search(r"(?:v=|youtu\.be/|/embed/|/shorts/)([A-Za-z0-9_-]{11})", url)
    if not m:
        raise HTTPException(status_code=400, detail="Could not detect a YouTube video id in that URL.")
    video_id = m.group(1)
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        # Optional Webshare proxy — YouTube blocks Render's datacenter IPs by
        # default. Set WEBSHARE_PROXY_USERNAME + WEBSHARE_PROXY_PASSWORD env vars
        # to route through Webshare's residential proxies (paid).
        proxy_config = None
        wh_user = os.getenv("WEBSHARE_PROXY_USERNAME")
        wh_pass = os.getenv("WEBSHARE_PROXY_PASSWORD")
        if wh_user and wh_pass:
            try:
                from youtube_transcript_api.proxies import WebshareProxyConfig
                proxy_config = WebshareProxyConfig(proxy_username=wh_user, proxy_password=wh_pass)
            except Exception as _e:
                print(f"[YouTube] Webshare proxy import failed: {_e}")

        # Support both library APIs:
        #   v0.6.x: YouTubeTranscriptApi.get_transcript(video_id) -> [{"text": ...}]
        #   v1.x:   YouTubeTranscriptApi().fetch(video_id) -> iterable of FetchedTranscriptSnippet
        if proxy_config is not None:
            transcript = YouTubeTranscriptApi(proxy_config=proxy_config).fetch(video_id)
            text = " ".join(getattr(c, "text", "") or (c.get("text", "") if isinstance(c, dict) else "") for c in transcript).strip()
        elif hasattr(YouTubeTranscriptApi, "get_transcript"):
            chunks = YouTubeTranscriptApi.get_transcript(video_id)
            text = " ".join((c.get("text", "") if isinstance(c, dict) else getattr(c, "text", "")) for c in chunks).strip()
        else:
            transcript = YouTubeTranscriptApi().fetch(video_id)
            text = " ".join(getattr(c, "text", "") or (c.get("text", "") if isinstance(c, dict) else "") for c in transcript).strip()
        if not text:
            raise HTTPException(status_code=422, detail="Transcript was empty.")
        return {"success": True, "video_id": video_id, "url": url, "text": text[:8000], "chars": len(text)}
    except HTTPException:
        raise
    except Exception as e:
        msg = str(e)
        # Detect YouTube's IP-block.
        blocked = ("blocking requests" in msg.lower()
                   or ("ip" in msg.lower() and "block" in msg.lower())
                   or "could not retrieve a transcript" in msg.lower())

        # If the transcript fetch failed, try the metadata fallback — title +
        # description scraped from the public page (works from any IP). This
        # gives us *something* the AI can use as source material instead of
        # leaving the teacher stuck.
        if blocked:
            fallback = await _youtube_metadata_fallback(video_id, url)
            if fallback:
                return {
                    "success": True,
                    "video_id": video_id,
                    "url": url,
                    "text": fallback["text"][:8000],
                    "chars": len(fallback["text"]),
                    "title": fallback["title"],
                    "note": "Transcript was blocked from our server's IP — using the video's title + description instead. For best results, paste the full transcript manually.",
                }

        raise HTTPException(
            status_code=502 if blocked else 500,
            detail=(
                "YouTube blocks transcript requests from our cloud server's IP, and we "
                "couldn't reach the video's public description either. Please open the "
                "video in your browser → click the ⋯ menu (or 'Show transcript' in the "
                "description) → copy the transcript → paste it into the textarea below."
                if blocked else f"YouTube transcript fetch failed: {msg}"
            ),
        )


# ── MCP Tools ─────────────────────────────────────────────────────────────────

@app.get("/mcp/tools")
async def list_mcp_tools():
    return {"tools": MCP_TOOLS}


@app.post("/mcp/tools/call")
async def call_mcp_tool(req: MCPToolCall):
    try:
        result = await execute_mcp_tool(req.tool_name, req.arguments)
        return {"success": True, "result": result}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health():
    return {"status": "ok", "tool": "reading-comprehension", "model": GROQ_MODEL}


# ── Serve frontend ────────────────────────────────────────────────────────────

frontend_dir = Path(__file__).parent.parent / "frontend" / "dist"

if frontend_dir.exists():
    from fastapi.responses import FileResponse

    @app.get("/")
    async def serve_index():
        return FileResponse(
            str(frontend_dir / "index.html"),
            headers={"Cache-Control": "no-store, no-cache, must-revalidate"},
        )

    @app.get("/assets/index.js")
    async def serve_js():
        return FileResponse(
            str(frontend_dir / "assets" / "index.js"),
            media_type="application/javascript",
            headers={"Cache-Control": "no-cache, must-revalidate"},
        )

    try:
        app.mount("/", StaticFiles(directory=str(frontend_dir), html=True), name="static")
    except Exception as e:
        print(f"[startup] Static files mount skipped: {e}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8003, reload=True)
